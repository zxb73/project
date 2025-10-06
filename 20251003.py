import os
import sys
import pandas as pd
import requests
import json
from datetime import datetime, timedelta
from docx import Document
import re
import traceback
from collections import defaultdict
import warnings
warnings.filterwarnings('ignore')

# PyQt5 imports
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                             QPushButton, QLabel, QTextEdit, QProgressBar, QListWidget, 
                             QLineEdit, QSplitter, QGroupBox, QMessageBox, QScrollArea,
                             QFrame, QTabWidget, QTableWidget, QTableWidgetItem, QHeaderView)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt5.QtGui import QFont, QPalette, QColor

class AnalysisWorker(QThread):
    """分析工作线程"""
    progress_updated = pyqtSignal(int, str)
    analysis_finished = pyqtSignal(str, str)  # 分析结果, 文件路径
    error_occurred = pyqtSignal(str)
    
    def __init__(self, sector_data, stock_data, prompt, api_key):
        super().__init__()
        self.sector_data = sector_data
        self.stock_data = stock_data
        self.prompt = prompt
        self.api_key = api_key
        self.base_url = "https://api.deepseek.com/v1/chat/completions"
    
    def run(self):
        try:
            self.progress_updated.emit(10, "开始数据分析...")
            
            # 准备数据上下文
            data_context = self.prepare_data_context()
            self.progress_updated.emit(30, "数据准备完成，开始调用AI分析...")
            
            # 调用API进行分析
            analysis_result = self.call_deepseek_api(self.prompt, data_context)
            self.progress_updated.emit(70, "AI分析完成，生成报告...")
            
            # 生成Word报告
            file_path = self.generate_word_report(analysis_result, self.prompt)
            self.progress_updated.emit(100, "分析完成！")
            
            self.analysis_finished.emit(analysis_result, file_path)
            
        except Exception as e:
            error_msg = f"分析过程中出错: {str(e)}\n{traceback.format_exc()}"
            self.error_occurred.emit(error_msg)
    
    def prepare_data_context(self):
        """准备数据上下文用于API调用"""
        context = "板块数据统计:\n"
        
        if self.sector_data is not None:
            context += f"- 数据时间范围: {self.sector_data['trade_date'].min()} 到 {self.sector_data['trade_date'].max()}\n"
            context += f"- 板块数量: {len(self.sector_data['代码'].unique())}\n"
            context += f"- 总记录数: {len(self.sector_data)}\n"
            
            # 添加关键统计信息
            numeric_columns = self.sector_data.select_dtypes(include=['number']).columns
            for col in numeric_columns[:5]:  # 只取前5个数值列
                if col not in ['代码', 'trade_date']:
                    context += f"- {col}均值: {self.sector_data[col].mean():.2f}\n"
        
        context += "\n个股数据统计:\n"
        if self.stock_data is not None:
            context += f"- 数据时间范围: {self.stock_data['trade_date'].min()} 到 {self.stock_data['trade_date'].max()}\n"
            context += f"- 股票数量: {len(self.stock_data['代码'].unique())}\n"
            context += f"- 总记录数: {len(self.stock_data)}\n"
            
            numeric_columns = self.stock_data.select_dtypes(include=['number']).columns
            for col in numeric_columns[:5]:
                if col not in ['代码', 'trade_date']:
                    context += f"- {col}均值: {self.stock_data[col].mean():.2f}\n"
        
        return context
    
    def call_deepseek_api(self, prompt, data_context):
        """调用DeepSeek API进行分析"""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        # 构建完整的提示词 - 添加限制条件
        full_prompt = f"""
        {prompt}
        
        数据背景:
        {data_context}
        
        请基于以上数据，按照以下要求进行分析：
        1. 分析板块数据，总结大盘规律和趋势
        2. 分析个股数据，识别有潜力的个股特征
        3. 结合大盘规律和个股特征，预测未来三天涨幅最大的10支股票
        4. 给出详细的投资建议和风险提示
        
        【重要限制条件】：
        - 预测名单必须使用个股数据里的数据，禁止自行发挥或添加数据中不存在的股票
        - 数据中的'代码'列是股票代码，'名称'列是股票名称
        - 只能基于提供的个股数据进行预测，不能引入外部知识
        
        请以专业的股票分析师角度进行回答，确保分析逻辑严谨。
        """
        
        payload = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": full_prompt
                }
            ],
            "temperature": 0.7,
            "max_tokens": 4000
        }
        
        try:
            response = requests.post(self.base_url, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
            result = response.json()
            return result['choices'][0]['message']['content']
        except Exception as e:
            return f"API调用失败: {str(e)}\n\n模拟分析结果:\n基于技术分析，推荐关注科技、新能源等板块的龙头股。建议结合当前市场趋势进行投资决策。"
    
    def generate_word_report(self, analysis_result, prompt):
        """生成Word报告"""
        doc = Document()
        
        # 标题
        title = doc.add_heading('股票分析报告', 0)
        title.alignment = 1
        
        # 报告信息
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"分析提示: {prompt}")
        doc.add_paragraph("")
        
        # 数据概况
        doc.add_heading('数据概况', level=1)
        if self.sector_data is not None:
            p = doc.add_paragraph()
            p.add_run('板块数据: ').bold = True
            p.add_run(f"共 {len(self.sector_data)} 条记录，涵盖 {len(self.sector_data['代码'].unique())} 个板块")
        
        if self.stock_data is not None:
            p = doc.add_paragraph()
            p.add_run('个股数据: ').bold = True
            p.add_run(f"共 {len(self.stock_data)} 条记录，涵盖 {len(self.stock_data['代码'].unique())} 支股票")
        
        # 分析结果
        doc.add_heading('分析结果', level=1)
        analysis_paragraphs = analysis_result.split('\n')
        for paragraph in analysis_paragraphs:
            if paragraph.strip():
                # 识别标题和重点内容
                if any(keyword in paragraph for keyword in ['推荐', '建议', '结论', '预测']):
                    p = doc.add_paragraph(paragraph)
                    p.style = 'List Bullet'
                else:
                    doc.add_paragraph(paragraph)
        
        # 风险提示
        doc.add_heading('风险提示', level=1)
        risk_note = """
        本报告基于历史数据和技术分析生成，仅供参考，不构成投资建议。
        股票市场存在风险，投资需谨慎。
        过去表现不代表未来收益，请结合自身风险承受能力做出投资决策。
        """
        doc.add_paragraph(risk_note)
        
        # 保存文件
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        filename = f"股票分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(desktop_path, filename)
        doc.save(filepath)
        
        return filepath


class ExcelReader:
    """Excel文件读取器，处理不同格式的Excel文件"""
    
    @staticmethod
    def read_excel_file(file_path):
        """读取Excel文件，尝试多种方式"""
        try:
            # 方法1: 尝试使用openpyxl引擎读取xlsx文件
            if file_path.endswith('.xlsx'):
                try:
                    df = pd.read_excel(file_path, engine='openpyxl')
                    if not df.empty and not df.isna().all().all():
                        return df, "openpyxl"
                except Exception as e:
                    print(f"openpyxl读取失败: {e}")
            
            # 方法2: 尝试使用xlrd引擎读取xls文件
            try:
                df = pd.read_excel(file_path, engine='xlrd')
                if not df.empty and not df.isna().all().all():
                    return df, "xlrd"
            except Exception as e:
                print(f"xlrd读取失败: {e}")
            
            # 方法3: 尝试自动检测引擎
            try:
                df = pd.read_excel(file_path)
                if not df.empty and not df.isna().all().all():
                    return df, "auto"
            except Exception as e:
                print(f"自动引擎读取失败: {e}")
            
            # 方法4: 尝试指定header=None
            try:
                df = pd.read_excel(file_path, header=None)
                if not df.empty and not df.isna().all().all():
                    return df, "no_header"
            except Exception as e:
                print(f"无表头读取失败: {e}")
            
            # 方法5: 尝试读取所有sheet
            try:
                xl = pd.ExcelFile(file_path)
                for sheet_name in xl.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    if not df.empty and not df.isna().all().all():
                        return df, f"sheet_{sheet_name}"
            except Exception as e:
                print(f"多sheet读取失败: {e}")
                
            return None, "all_failed"
            
        except Exception as e:
            print(f"读取Excel文件异常: {e}")
            return None, "exception"


class StockAnalyzerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.sector_data = None
        self.stock_data = None
        # 从环境变量读取 DeepSeek API 密钥，便于分发时不在代码中明文保存
        # 若未设置环境变量，程序启动后会在界面中提示输入密钥
        self.api_key = os.environ.get("DEEPSEEK_API_KEY", "")
        self.analysis_worker = None
        
        self.init_ui()
        
    def init_ui(self):
        """初始化UI"""
        self.setWindowTitle("股票数据分析系统 - 专业版")
        self.setGeometry(100, 100, 1200, 800)
        
        # 设置样式 - 苹果风格
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f7fa;
            }
            QGroupBox {
                font-weight: bold;
                border: 1px solid #d1d5db;
                border-radius: 10px;
                margin-top: 1ex;
                padding-top: 10px;
                background-color: white;
                color: #1f2937;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
                color: #374151;
                font-size: 14px;
            }
            QPushButton {
                background-color: #3b82f6;
                border: none;
                color: white;
                padding: 10px 20px;
                text-align: center;
                text-decoration: none;
                font-size: 14px;
                margin: 4px 2px;
                border-radius: 8px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2563eb;
            }
            QPushButton:pressed {
                background-color: #1d4ed8;
            }
            QPushButton:disabled {
                background-color: #9ca3af;
                color: #d1d5db;
            }
            QPushButton#danger {
                background-color: #ef4444;
            }
            QPushButton#danger:hover {
                background-color: #dc2626;
            }
            QPushButton#danger:pressed {
                background-color: #b91c1c;
            }
            QTextEdit, QListWidget, QLineEdit, QTableWidget {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                padding: 8px;
                background-color: white;
                color: #1f2937;
                font-size: 13px;
                selection-background-color: #3b82f6;
            }
            QProgressBar {
                border: 1px solid #d1d5db;
                border-radius: 8px;
                text-align: center;
                background-color: #e5e7eb;
                color: #374151;
                height: 20px;
            }
            QProgressBar::chunk {
                background-color: #10b981;
                border-radius: 7px;
            }
            QLabel {
                color: #374151;
                background-color: transparent;
            }
            QTabWidget::pane {
                border: 1px solid #d1d5db;
                background-color: white;
                border-radius: 10px;
                margin-top: 10px;
            }
            QTabWidget::tab-bar {
                alignment: center;
            }
            QTabBar {
                background-color: #f3f4f6;
                border-radius: 8px;
                margin: 5px;
            }
            QTabBar::tab {
                background-color: #f3f4f6;
                color: #6b7280;
                padding: 12px 24px;
                margin-right: 2px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
                font-size: 14px;
                font-weight: bold;
                min-width: 140px;
            }
            QTabBar::tab:first {
                border-top-left-radius: 8px;
                border-bottom-left-radius: 8px;
            }
            QTabBar::tab:last {
                border-top-right-radius: 8px;
                border-bottom-right-radius: 8px;
                margin-right: 0;
            }
            QTabBar::tab:selected {
                background-color: #3b82f6;
                color: white;
                font-weight: bold;
            }
            QTabBar::tab:hover:!selected {
                background-color: #e5e7eb;
                color: #374151;
            }
            QListWidget::item {
                padding: 8px;
                border-bottom: 1px solid #f3f4f6;
            }
            QListWidget::item:selected {
                background-color: #3b82f6;
                color: white;
                border-radius: 5px;
            }
            QScrollArea {
                border: none;
                background-color: transparent;
            }
            QScrollBar:vertical {
                border: none;
                background-color: #f3f4f6;
                width: 12px;
                margin: 0px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #d1d5db;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #9ca3af;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                border: none;
                background: none;
            }
        """)
        
        # 创建中央窗口部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 创建主布局
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(15)
        
        # 标题
        title_label = QLabel("股票数据分析系统")
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setFont(QFont("Arial", 20, QFont.Bold))
        title_label.setStyleSheet("color: #1f2937; padding: 15px; background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #3b82f6, stop:1 #8b5cf6); color: white; border-radius: 10px;")
        main_layout.addWidget(title_label)
        
        # 创建选项卡 - 设置为顶部横向布局
        tab_widget = QTabWidget()
        tab_widget.setTabPosition(QTabWidget.North)  # 将选项卡放在顶部
        main_layout.addWidget(tab_widget)
        
        # 数据加载选项卡
        data_tab = QWidget()
        tab_widget.addTab(data_tab, "📊 数据加载")
        self.setup_data_tab(data_tab)
        
        # 分析结果选项卡
        analysis_tab = QWidget()
        tab_widget.addTab(analysis_tab, "🔍 分析结果")
        self.setup_analysis_tab(analysis_tab)
        
        # 日志选项卡
        log_tab = QWidget()
        tab_widget.addTab(log_tab, "📝 执行日志")
        self.setup_log_tab(log_tab)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)
        
        # 状态标签
        self.status_label = QLabel("就绪")
        self.status_label.setStyleSheet("color: #6b7280; font-style: italic; padding: 8px; background-color: transparent;")
        main_layout.addWidget(self.status_label)
        
    def setup_data_tab(self, parent):
        """设置数据加载选项卡"""
        layout = QVBoxLayout(parent)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(15)
        
        # 文件选择区域
        file_group = QGroupBox("文件选择")
        file_layout = QVBoxLayout(file_group)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        self.select_folder_btn = QPushButton("📁 选择文件夹")
        self.select_folder_btn.clicked.connect(self.load_folder)
        
        self.select_files_btn = QPushButton("📄 选择文件")
        self.select_files_btn.clicked.connect(self.load_files)
        
        self.clear_data_btn = QPushButton("🗑️ 清空数据")
        self.clear_data_btn.clicked.connect(self.clear_data)
        self.clear_data_btn.setObjectName("danger")
        
        button_layout.addWidget(self.select_folder_btn)
        button_layout.addWidget(self.select_files_btn)
        button_layout.addWidget(self.clear_data_btn)
        button_layout.addStretch()
        
        file_layout.addLayout(button_layout)
        
        # 文件列表
        file_list_label = QLabel("已选文件:")
        file_list_label.setStyleSheet("font-weight: bold; color: #374151;")
        file_layout.addWidget(file_list_label)
        
        self.file_list_widget = QListWidget()
        file_layout.addWidget(self.file_list_widget)
        
        layout.addWidget(file_group)
        
        # 数据预览区域
        preview_group = QGroupBox("数据预览")
        preview_layout = QVBoxLayout(preview_group)
        
        # 创建滚动区域用于预览
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        
        preview_widget = QWidget()
        preview_scroll_layout = QVBoxLayout(preview_widget)
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setStyleSheet("font-family: 'SF Mono', 'Monaco', 'Courier New', monospace; font-size: 11px; background-color: white; color: #1f2937;")
        preview_scroll_layout.addWidget(self.preview_text)
        
        scroll_area.setWidget(preview_widget)
        preview_layout.addWidget(scroll_area)
        
        layout.addWidget(preview_group)
        
    def setup_analysis_tab(self, parent):
        """设置分析结果选项卡"""
        layout = QVBoxLayout(parent)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(15)
        
        # 分析设置区域
        settings_group = QGroupBox("分析设置")
        settings_layout = QVBoxLayout(settings_group)
        
        # API密钥设置
        api_layout = QHBoxLayout()
        api_label = QLabel("API密钥:")
        api_label.setStyleSheet("font-weight: bold; color: #374151;")
        api_layout.addWidget(api_label)
        self.api_key_edit = QLineEdit(self.api_key)
        self.api_key_edit.setPlaceholderText("请输入DeepSeek API密钥")
        self.api_key_edit.setEchoMode(QLineEdit.Password)  # 设置为密码模式，显示星号
        api_layout.addWidget(self.api_key_edit)
        settings_layout.addLayout(api_layout)
        
        # 提示词设置
        prompt_layout = QVBoxLayout()
        prompt_label = QLabel("分析提示词:")
        prompt_label.setStyleSheet("font-weight: bold; color: #374151;")
        prompt_layout.addWidget(prompt_label)
        self.prompt_edit = QTextEdit()
        self.prompt_edit.setMaximumHeight(120)
        self.prompt_edit.setPlainText("请分析以下股票数据，预测未来三天涨幅最大的10支股票")
        self.prompt_edit.setPlaceholderText("请输入分析提示词...")
        prompt_layout.addWidget(self.prompt_edit)
        settings_layout.addLayout(prompt_layout)
        
        # 分析按钮
        self.analyze_btn = QPushButton("🚀 开始分析")
        self.analyze_btn.clicked.connect(self.analyze_data)
        settings_layout.addWidget(self.analyze_btn)
        
        layout.addWidget(settings_group)
        
        # 分析结果区域
        result_group = QGroupBox("分析结果")
        result_layout = QVBoxLayout(result_group)
        
        # 创建滚动区域用于结果显示
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        
        result_widget = QWidget()
        result_scroll_layout = QVBoxLayout(result_widget)
        
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        self.result_text.setStyleSheet("font-size: 12px; line-height: 1.5; background-color: white; color: #1f2937;")
        result_scroll_layout.addWidget(self.result_text)
        
        scroll_area.setWidget(result_widget)
        result_layout.addWidget(scroll_area)
        
        layout.addWidget(result_group)
        
    def setup_log_tab(self, parent):
        """设置执行日志选项卡"""
        layout = QVBoxLayout(parent)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(15)
        
        # 日志区域
        log_group = QGroupBox("执行日志")
        log_layout = QVBoxLayout(log_group)
        
        # 日志控制按钮
        log_control_layout = QHBoxLayout()
        self.clear_log_btn = QPushButton("🗑️ 清空日志")
        self.clear_log_btn.clicked.connect(self.clear_log)
        self.clear_log_btn.setObjectName("danger")
        
        log_control_layout.addWidget(self.clear_log_btn)
        log_control_layout.addStretch()
        log_layout.addLayout(log_control_layout)
        
        # 创建滚动区域用于日志显示
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        
        log_widget = QWidget()
        log_scroll_layout = QVBoxLayout(log_widget)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setStyleSheet("font-family: 'SF Mono', 'Monaco', 'Courier New', monospace; font-size: 11px; background-color: #1f2937; color: #e5e7eb;")
        log_scroll_layout.addWidget(self.log_text)
        
        scroll_area.setWidget(log_widget)
        log_layout.addWidget(scroll_area)
        
        layout.addWidget(log_group)
        
    def log_message(self, message):
        """记录日志消息"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.log_text.append(f"[{timestamp}] {message}")
        # 自动滚动到底部
        self.log_text.verticalScrollBar().setValue(
            self.log_text.verticalScrollBar().maximum()
        )
        
    def clear_log(self):
        """清空日志"""
        self.log_text.clear()
        
    def load_folder(self):
        """加载文件夹"""
        from PyQt5.QtWidgets import QFileDialog
        folder_path = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if folder_path:
            self.process_files(folder_path)
    
    def load_files(self):
        """加载文件"""
        from PyQt5.QtWidgets import QFileDialog
        file_paths, _ = QFileDialog.getOpenFileNames(
            self, "选择文件", "", "Excel files (*.xls *.xlsx);;All files (*.*)"
        )
        if file_paths:
            self.process_files(file_paths)
    
    def extract_date_from_filename(self, filename):
        """从文件名中提取日期"""
        match = re.search(r'(\d{8})', filename)
        if match:
            date_str = match.group(1)
            return datetime.strptime(date_str, "%Y%m%d")
        return None
    
    def clean_dataframe(self, df):
        """清洗数据框，剔除超过一半为NaN的行"""
        if df is None or df.empty:
            return df
            
        original_rows = len(df)
        # 计算每行NaN值的比例
        nan_ratio = df.isna().sum(axis=1) / df.shape[1]
        # 保留NaN比例小于0.5的行
        cleaned_df = df[nan_ratio < 0.5].copy()
        removed_rows = original_rows - len(cleaned_df)
        
        if removed_rows > 0:
            self.log_message(f"⚠️ 数据清洗: 移除了 {removed_rows} 行(超过50%为空值的数据)")
            
        return cleaned_df
    
    def read_excel_with_retry(self, file_path):
        """使用多种方法尝试读取Excel文件"""
        df, method = ExcelReader.read_excel_file(file_path)
        
        if df is not None:
            # 数据清洗 - 剔除超过一半为NaN的行
            df = self.clean_dataframe(df)
            
            self.log_message(f"✅ 成功读取文件: {os.path.basename(file_path)} (方法: {method})")
            
            # 检查数据质量
            if not df.empty:
                nan_percentage = df.isna().sum().sum() / (df.shape[0] * df.shape[1])
                if nan_percentage > 0.5:  # 如果50%以上是NaN，发出警告
                    self.log_message(f"⚠️ 文件 {os.path.basename(file_path)} 包含较多空值 ({nan_percentage:.1%})")
                
                # 显示数据基本信息
                self.log_message(f"  数据形状: {df.shape}, 列名: {list(df.columns)}")
            else:
                self.log_message(f"⚠️ 文件 {os.path.basename(file_path)} 清洗后为空")
            
            return df
        else:
            self.log_message(f"❌ 无法读取文件: {os.path.basename(file_path)} (所有方法都失败)")
            return None
    
    def process_files(self, file_source):
        """处理文件数据"""
        all_files = []
        
        if isinstance(file_source, str):  # 文件夹路径
            for root, dirs, files in os.walk(file_source):
                for file in files:
                    if file.endswith(('.xls', '.xlsx')):
                        all_files.append(os.path.join(root, file))
        else:  # 文件列表
            all_files = file_source
        
        if not all_files:
            QMessageBox.warning(self, "警告", "未找到Excel文件")
            return
        
        # 清空文件列表显示
        self.file_list_widget.clear()
        for file in all_files:
            self.file_list_widget.addItem(file)
        
        # 分类处理文件
        sector_files = []
        stock_files = []
        
        for file in all_files:
            filename = os.path.basename(file)
            if '板块' in filename:
                sector_files.append(file)
            else:
                stock_files.append(file)
        
        # 读取并处理板块数据
        sector_data_list = []
        for file in sector_files:
            try:
                date = self.extract_date_from_filename(os.path.basename(file))
                df = self.read_excel_with_retry(file)
                
                if df is not None and not df.empty:
                    df['trade_date'] = date
                    df['file_type'] = 'sector'
                    sector_data_list.append(df)
            except Exception as e:
                error_msg = f"❌ 处理板块文件 {file} 时出错: {e}"
                self.log_message(error_msg)
        
        # 读取并处理个股数据
        stock_data_list = []
        for file in stock_files:
            try:
                date = self.extract_date_from_filename(os.path.basename(file))
                df = self.read_excel_with_retry(file)
                
                if df is not None and not df.empty:
                    df['trade_date'] = date
                    df['file_type'] = 'stock'
                    stock_data_list.append(df)
            except Exception as e:
                error_msg = f"❌ 处理个股文件 {file} 时出错: {e}"
                self.log_message(error_msg)
        
        # 合并数据
        if sector_data_list:
            self.sector_data = pd.concat(sector_data_list, ignore_index=True)
            # 按日期和代码归类
            self.sector_data = self.sector_data.sort_values(['trade_date', '代码'])
            self.log_message(f"✅ 板块数据合并完成: {len(self.sector_data)} 条记录")
        
        if stock_data_list:
            self.stock_data = pd.concat(stock_data_list, ignore_index=True)
            # 按日期和代码归类
            self.stock_data = self.stock_data.sort_values(['trade_date', '代码'])
            self.log_message(f"✅ 个股数据合并完成: {len(self.stock_data)} 条记录")
        
        # 显示预览
        self.show_preview()
        
        success_msg = f"✅ 数据加载完成! 板块文件: {len(sector_files)}个, 个股文件: {len(stock_files)}个"
        self.log_message(success_msg)
        QMessageBox.information(self, "成功", success_msg)
    
    def show_preview(self):
        """显示数据预览"""
        self.preview_text.clear()
        
        if self.sector_data is not None and not self.sector_data.empty:
            self.preview_text.append("=== 板块数据预览 (前10行) ===")
            preview_df = self.sector_data.head(10)
            # 确保显示所有列
            pd.set_option('display.max_columns', None)
            self.preview_text.append(preview_df.to_string())
            self.preview_text.append("\n")
            
            # 显示数据基本信息
            self.preview_text.append(f"板块数据基本信息:")
            self.preview_text.append(f"- 总行数: {len(self.sector_data)}")
            self.preview_text.append(f"- 总列数: {len(self.sector_data.columns)}")
            self.preview_text.append(f"- 列名: {list(self.sector_data.columns)}")
            self.preview_text.append(f"- 数据日期范围: {self.sector_data['trade_date'].min()} 到 {self.sector_data['trade_date'].max()}")
            self.preview_text.append("\n")
        
        if self.stock_data is not None and not self.stock_data.empty:
            self.preview_text.append("=== 个股数据预览 (前10行) ===")
            preview_df = self.stock_data.head(10)
            self.preview_text.append(preview_df.to_string())
            
            # 显示数据基本信息
            self.preview_text.append(f"\n个股数据基本信息:")
            self.preview_text.append(f"- 总行数: {len(self.stock_data)}")
            self.preview_text.append(f"- 总列数: {len(self.stock_data.columns)}")
            self.preview_text.append(f"- 列名: {list(self.stock_data.columns)}")
            self.preview_text.append(f"- 数据日期范围: {self.stock_data['trade_date'].min()} 到 {self.stock_data['trade_date'].max()}")
    
    def analyze_data(self):
        """分析数据"""
        if self.sector_data is None and self.stock_data is None:
            QMessageBox.warning(self, "警告", "请先加载数据")
            return
        
        prompt = self.prompt_edit.toPlainText()
        if not prompt.strip():
            QMessageBox.warning(self, "警告", "请输入分析提示词")
            return
        
        # 更新API密钥
        self.api_key = self.api_key_edit.text().strip()
        if not self.api_key:
            QMessageBox.warning(self, "警告", "请输入API密钥")
            return
        
        # 禁用分析按钮
        self.analyze_btn.setEnabled(False)
        self.analyze_btn.setText("分析中...")
        
        # 显示进度条
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        
        # 创建工作线程
        self.analysis_worker = AnalysisWorker(
            self.sector_data, self.stock_data, prompt, self.api_key
        )
        self.analysis_worker.progress_updated.connect(self.update_progress)
        self.analysis_worker.analysis_finished.connect(self.analysis_completed)
        self.analysis_worker.error_occurred.connect(self.analysis_error)
        
        # 启动分析
        self.analysis_worker.start()
        
        self.log_message("🚀 开始数据分析...")
    
    def update_progress(self, value, message):
        """更新进度"""
        self.progress_bar.setValue(value)
        self.status_label.setText(message)
        self.log_message(f"📊 进度 {value}%: {message}")
    
    def analysis_completed(self, result, file_path):
        """分析完成"""
        self.progress_bar.setValue(100)
        self.status_label.setText("✅ 分析完成！")
        self.analyze_btn.setEnabled(True)
        self.analyze_btn.setText("🚀 开始分析")
        
        # 显示分析结果
        self.result_text.setPlainText(result)
        
        # 记录完成信息
        self.log_message(f"✅ 分析完成！报告已保存到: {file_path}")
        
        # 显示成功消息
        QMessageBox.information(self, "成功", f"分析完成！\n报告已保存到:\n{file_path}")
    
    def analysis_error(self, error_msg):
        """分析出错"""
        self.progress_bar.setVisible(False)
        self.status_label.setText("❌ 分析出错")
        self.analyze_btn.setEnabled(True)
        self.analyze_btn.setText("🚀 开始分析")
        
        # 记录错误
        self.log_message(f"❌ 分析错误: {error_msg}")
        
        # 显示错误消息
        QMessageBox.critical(self, "错误", f"分析过程中出错:\n{error_msg}")
    
    def clear_data(self):
        """清空数据"""
        self.sector_data = None
        self.stock_data = None
        self.file_list_widget.clear()
        self.preview_text.clear()
        self.result_text.clear()
        self.log_message("🗑️ 数据已清空")
        QMessageBox.information(self, "成功", "数据已清空")

def main():
    """主函数"""
    app = QApplication(sys.argv)
    
    # 设置应用程序样式
    app.setStyle('Fusion')
    
    # 设置应用程序字体
    font = QFont("Microsoft YaHei UI", 10)  # 使用微软雅黑字体，更好的中文支持
    app.setFont(font)
    
    # 创建并显示主窗口
    window = StockAnalyzerApp()
    window.show()
    
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()