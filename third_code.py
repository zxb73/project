import os
import sys
import pandas as pd
import numpy as np
from openai import OpenAI
from datetime import datetime, timedelta
import glob
import re
from docx import Document
import warnings
warnings.filterwarnings('ignore')

# PyQt5 imports
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
                            QLabel, QLineEdit, QPushButton, QProgressBar, QTextEdit,
                            QGroupBox, QRadioButton, QButtonGroup, QFileDialog, QMessageBox,
                            QTableWidget, QTableWidgetItem, QHeaderView, QTabWidget, QCheckBox)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont

class AnalysisThread(QThread):
    """分析线程"""
    progress_updated = pyqtSignal(int, str, str)
    log_message = pyqtSignal(str)
    analysis_finished = pyqtSignal(dict)
    analysis_error = pyqtSignal(str)
    
    def __init__(self, data_folder, analysis_type, predict_days, api_key):
        super().__init__()
        self.data_folder = data_folder
        self.analysis_type = analysis_type
        self.predict_days = predict_days
        self.api_key = api_key
        self.stock_data = {}
        self.market_data = []
    
    def run(self):
        try:
            self.log_message.emit("🚀 开始股票数据分析...")
            
            # 步骤1: 扫描和分类文件
            self.progress_updated.emit(10, "扫描数据文件...", "")
            market_files, stock_files = self.scan_data_files()
            self.log_message.emit(f"找到 {len(market_files)} 个大盘文件, {len(stock_files)} 个个股文件")
            
            if not market_files and not stock_files:
                self.analysis_error.emit("未找到任何Excel文件，请检查文件夹路径和文件格式")
                return
            
            # 步骤2: 处理大盘数据
            self.progress_updated.emit(30, "处理大盘数据...", "")
            market_data = self.process_market_data(market_files)
            self.log_message.emit(f"成功处理 {len(market_data)} 个大盘数据文件")
            
            # 步骤3: 处理个股数据
            self.progress_updated.emit(50, "处理个股数据...", "")
            stock_data = self.process_stock_data(stock_files)
            self.log_message.emit(f"成功处理 {len(stock_data)} 只股票数据")
            
            # 检查是否有足够的数据
            if not stock_data:
                self.log_message.emit("⚠️ 未能成功处理任何个股数据，尝试生成基础报告")
                # 即使没有个股数据，也生成基础报告
                self.generate_basic_report(market_files, stock_files, market_data)
                return
            
            # 步骤4: 计算股票收益（带排错功能）
            self.progress_updated.emit(70, "计算股票收益率...", "")
            returns_data = self.calculate_stock_returns_with_fallback(stock_data)
            self.log_message.emit(f"成功计算 {len(returns_data)} 只股票的收益率")
            
            if not returns_data:
                self.log_message.emit("⚠️ 无法计算股票收益率，生成基础分析报告")
                self.generate_basic_analysis_report(market_data, stock_data, market_files, stock_files)
                return
            
            # 获取涨幅前10的股票
            sorted_returns = sorted(returns_data.items(), key=lambda x: x[1]['total_return'], reverse=True)
            top_10_stocks = dict(list(sorted_returns)[:min(10, len(returns_data))])
            
            self.log_message.emit(f"找到 {len(top_10_stocks)} 只表现优秀的股票")
            
            # 步骤5: 使用DeepSeek分析
            market_analysis = "基于历史数据的分析"
            stock_analysis = {}
            
            if self.api_key and self.api_key != "your-api-key-here":
                try:
                    self.progress_updated.emit(80, "使用DeepSeek分析大盘趋势...", "")
                    market_analysis = self.analyze_market_with_deepseek(market_data)
                    self.log_message.emit("大盘趋势分析完成")
                    
                    self.progress_updated.emit(85, "使用DeepSeek分析个股...", "")
                    stock_analysis = self.analyze_stocks_with_deepseek(top_10_stocks, stock_data, market_analysis)
                    self.log_message.emit("个股分析完成")
                except Exception as e:
                    self.log_message.emit(f"⚠️ DeepSeek分析失败，使用基础分析: {str(e)}")
                    market_analysis = self.generate_basic_market_analysis(market_data)
                    stock_analysis = self.generate_basic_stock_analysis(top_10_stocks)
            else:
                self.log_message.emit("⚠️ 未配置有效的API密钥，使用基础分析")
                market_analysis = self.generate_basic_market_analysis(market_data)
                stock_analysis = self.generate_basic_stock_analysis(top_10_stocks)
            
            # 步骤6: 生成Word报告
            self.progress_updated.emit(90, "生成分析报告...", "")
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            output_path = os.path.join(desktop_path, f"股票分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            
            self.create_word_report(top_10_stocks, market_analysis, stock_analysis, output_path)
            
            # 完成
            self.progress_updated.emit(100, "分析完成！", "")
            self.log_message.emit("✅ 股票分析完成！")
            self.log_message.emit(f"📄 报告已保存到: {output_path}")
            
            # 返回结果
            result = {
                'market_files': len(market_files),
                'stock_files': len(stock_files),
                'market_data': len(market_data),
                'stock_data': len(stock_data),
                'top_stocks': top_10_stocks,
                'stock_analysis': stock_analysis,
                'report_path': output_path,
                'file_preview': self.get_file_preview(market_files + stock_files)
            }
            self.analysis_finished.emit(result)
            
        except Exception as e:
            self.analysis_error.emit(str(e))
    
    def scan_data_files(self):
        """扫描数据文件"""
        market_files = []
        stock_files = []
        
        for pattern in ["*.xls", "*.xlsx"]:
            search_pattern = os.path.join(self.data_folder, "**", pattern)
            for file_path in glob.glob(search_pattern, recursive=True):
                filename = os.path.basename(file_path)
                if '板块' in filename:
                    market_files.append(file_path)
                else:
                    stock_files.append(file_path)
        
        return market_files, stock_files
    
    def get_file_preview(self, file_paths):
        """获取文件预览信息"""
        preview_info = {}
        for file_path in file_paths[:3]:  # 只预览前3个文件
            try:
                df = self.read_excel_file(file_path)
                if df is not None:
                    preview_info[os.path.basename(file_path)] = {
                        'shape': df.shape,
                        'columns': df.columns.tolist(),
                        'first_5_rows': df.head().to_dict('records')
                    }
            except Exception as e:
                preview_info[os.path.basename(file_path)] = f"读取失败: {str(e)}"
        
        return preview_info
    
    def extract_date_from_filename(self, filename):
        """从文件名中提取日期"""
        match = re.search(r'(\d{8})', filename)
        if match:
            date_str = match.group(1)
            return f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
        return None
    
    def read_excel_file(self, file_path):
        """读取Excel文件（兼容低版本）"""
        try:
            self.log_message.emit(f"尝试读取文件: {os.path.basename(file_path)}")
            
            # 尝试多种读取方式
            engines = [None, 'xlrd', 'openpyxl']
            for engine in engines:
                try:
                    df = pd.read_excel(file_path, engine=engine)
                    if not df.empty:
                        self.log_message.emit(f"✅ 使用引擎 {engine} 成功读取 {os.path.basename(file_path)}")
                        self.log_message.emit(f"   数据形状: {df.shape}, 列名: {list(df.columns)}")
                        return df
                except Exception as e:
                    self.log_message.emit(f"   引擎 {engine} 失败: {str(e)}")
                    continue
            
            # 如果常规方法失败，尝试其他方式
            try:
                import xlwings as xw
                self.log_message.emit(f"尝试使用xlwings读取: {os.path.basename(file_path)}")
                app = xw.App(visible=False)
                wb = app.books.open(file_path)
                sheet = wb.sheets[0]
                data_range = sheet.used_range
                df = data_range.options(pd.DataFrame, index=False, header=True).value
                wb.close()
                app.quit()
                if df is not None and not df.empty:
                    self.log_message.emit(f"✅ 使用xlwings成功读取 {os.path.basename(file_path)}")
                    return df
            except Exception as e:
                self.log_message.emit(f"   xlwings读取失败: {str(e)}")
                
            return None
                
        except Exception as e:
            self.log_message.emit(f"❌ 读取文件失败 {os.path.basename(file_path)}: {str(e)}")
            return None
    
    def process_market_data(self, file_paths):
        """处理大盘数据"""
        market_data = []
        
        for file_path in file_paths:
            df = self.read_excel_file(file_path)
            if df is not None and not df.empty:
                # 添加日期信息
                date_str = self.extract_date_from_filename(os.path.basename(file_path))
                if date_str:
                    df['统计日期'] = date_str
                    market_data.append(df)
                    self.log_message.emit(f"✅ 成功处理大盘文件: {os.path.basename(file_path)}")
                else:
                    self.log_message.emit(f"⚠️ 无法从文件名提取日期: {os.path.basename(file_path)}")
        
        return market_data
    
    def process_stock_data(self, file_paths):
        """处理个股数据"""
        stock_data = {}
        
        for file_path in file_paths:
            df = self.read_excel_file(file_path)
            if df is not None and not df.empty:
                date_str = self.extract_date_from_filename(os.path.basename(file_path))
                if date_str:
                    df['统计日期'] = date_str
                    
                    # 查找股票代码列
                    code_columns = ['代码', '股票代码', 'symbol', 'code', '证券代码', '股票编码']
                    code_col = None
                    for col in df.columns:
                        col_str = str(col).lower()
                        if any(code_col in col_str for code_col in ['代码', 'code', 'symbol']):
                            code_col = col
                            break
                    
                    if code_col:
                        self.log_message.emit(f"找到代码列: {code_col}")
                        for _, row in df.iterrows():
                            try:
                                stock_code = str(row[code_col]).strip()
                                if stock_code and stock_code != 'nan':
                                    if stock_code not in stock_data:
                                        stock_data[stock_code] = []
                                    stock_data[stock_code].append(row.to_dict())
                            except Exception as e:
                                self.log_message.emit(f"处理行数据失败: {str(e)}")
                        self.log_message.emit(f"✅ 成功处理个股文件: {os.path.basename(file_path)}")
                    else:
                        self.log_message.emit(f"⚠️ 未找到股票代码列，文件: {os.path.basename(file_path)}")
                        self.log_message.emit(f"   可用列: {list(df.columns)}")
        
        # 转换为DataFrame
        processed_data = {}
        for code, records in stock_data.items():
            if records:  # 确保有记录
                try:
                    df = pd.DataFrame(records)
                    # 确保统计日期列存在
                    if '统计日期' in df.columns:
                        df['统计日期'] = pd.to_datetime(df['统计日期'], errors='coerce')
                        df = df.dropna(subset=['统计日期'])
                        df = df.sort_values('统计日期')
                        processed_data[code] = df
                        self.log_message.emit(f"✅ 成功整理股票 {code} 的数据，共 {len(df)} 条记录")
                except Exception as e:
                    self.log_message.emit(f"❌ 整理股票 {code} 数据失败: {str(e)}")
        
        return processed_data
    
    def calculate_stock_returns_with_fallback(self, stock_data):
        """计算股票收益率（带排错功能）"""
        returns_data = {}
        price_columns_tried = set()  # 记录尝试过的价格列
        
        for code, df in stock_data.items():
            if len(df) < 2:
                self.log_message.emit(f"⚠️ 股票 {code} 数据不足，跳过")
                continue
            
            # 查找价格列
            price_cols = ['收盘', '收盘价', 'close', 'Close', '价格', '最新价', '现价', '当前价']
            price_col = None
            
            for col in df.columns:
                col_str = str(col).lower()
                for price_keyword in ['收盘', 'close', '价格', '价', 'last']:
                    if price_keyword in col_str:
                        price_col = col
                        price_columns_tried.add(col)
                        break
                if price_col:
                    break
            
            if not price_col:
                self.log_message.emit(f"⚠️ 股票 {code} 未找到价格列，跳过。可用列: {list(df.columns)}")
                continue
            
            try:
                # 尝试转换为数值类型
                df[price_col] = pd.to_numeric(df[price_col], errors='coerce')
                df_sorted = df.sort_values('统计日期')
                df_sorted = df_sorted.dropna(subset=[price_col])
                
                if len(df_sorted) < 2:
                    self.log_message.emit(f"⚠️ 股票 {code} 有效价格数据不足，跳过")
                    continue
                
                start_price = df_sorted[price_col].iloc[0]
                end_price = df_sorted[price_col].iloc[-1]
                
                if start_price <= 0 or pd.isna(start_price) or pd.isna(end_price):
                    self.log_message.emit(f"⚠️ 股票 {code} 价格数据无效，跳过")
                    continue
                
                total_return = (end_price - start_price) / start_price * 100
                
                returns_data[code] = {
                    'total_return': total_return,
                    'start_price': start_price,
                    'end_price': end_price,
                    'data_points': len(df_sorted),
                    'price_column': price_col
                }
                self.log_message.emit(f"✅ 计算股票 {code} 收益率: {total_return:.2f}% (使用列: {price_col})")
                
            except Exception as e:
                self.log_message.emit(f"⚠️ 计算股票 {code} 收益率失败: {str(e)}，跳过")
                continue
        
        # 如果没有找到任何收益率数据，尝试备选方案
        if not returns_data and stock_data:
            self.log_message.emit("⚠️ 未找到标准价格列，尝试使用第一列数值数据")
            returns_data = self.calculate_returns_using_first_numeric(stock_data)
        
        return returns_data
    
    def calculate_returns_using_first_numeric(self, stock_data):
        """使用第一列数值数据计算收益率（备选方案）"""
        returns_data = {}
        
        for code, df in stock_data.items():
            if len(df) < 2:
                continue
            
            # 查找第一个数值列
            numeric_col = None
            for col in df.columns:
                try:
                    # 尝试转换为数值
                    numeric_series = pd.to_numeric(df[col], errors='coerce')
                    if numeric_series.notna().sum() >= 2:  # 至少有2个有效数值
                        numeric_col = col
                        break
                except:
                    continue
            
            if numeric_col:
                try:
                    df_sorted = df.sort_values('统计日期')
                    numeric_values = pd.to_numeric(df_sorted[numeric_col], errors='coerce')
                    valid_data = df_sorted[numeric_values.notna()]
                    
                    if len(valid_data) >= 2:
                        start_val = numeric_values.iloc[0]
                        end_val = numeric_values.iloc[-1]
                        
                        if start_val > 0:
                            total_return = (end_val - start_val) / start_val * 100
                            
                            returns_data[code] = {
                                'total_return': total_return,
                                'start_price': start_val,
                                'end_price': end_val,
                                'data_points': len(valid_data),
                                'price_column': f"{numeric_col}(备选)"
                            }
                            self.log_message.emit(f"✅ 使用备选列计算股票 {code} 收益率: {total_return:.2f}% (使用列: {numeric_col})")
                except Exception as e:
                    self.log_message.emit(f"⚠️ 备选方案计算股票 {code} 收益率失败: {str(e)}")
        
        return returns_data
    
    def generate_basic_report(self, market_files, stock_files, market_data):
        """生成基础报告（当没有个股数据时）"""
        try:
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            output_path = os.path.join(desktop_path, f"股票分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            
            doc = Document()
            doc.add_heading('股票数据扫描报告', 0)
            
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"扫描文件夹: {self.data_folder}")
            doc.add_paragraph("")
            
            doc.add_heading('扫描结果', level=1)
            doc.add_paragraph(f"找到大盘文件: {len(market_files)} 个")
            doc.add_paragraph(f"找到个股文件: {len(stock_files)} 个")
            doc.add_paragraph(f"成功处理大盘数据: {len(market_data)} 个")
            doc.add_paragraph("")
            
            doc.add_heading('问题分析', level=1)
            doc.add_paragraph("未能成功分析个股数据，可能的原因:")
            doc.add_paragraph("1. 文件格式不兼容")
            doc.add_paragraph("2. 未找到股票代码列")
            doc.add_paragraph("3. 数据列名不标准")
            doc.add_paragraph("4. 文件内容为空或格式错误")
            doc.add_paragraph("")
            
            doc.add_heading('建议', level=1)
            doc.add_paragraph("1. 检查文件是否为标准Excel格式")
            doc.add_paragraph("2. 确认文件包含股票代码和价格信息")
            doc.add_paragraph("3. 查看详细日志了解具体错误")
            
            doc.save(output_path)
            
            self.log_message.emit(f"📄 生成基础报告: {output_path}")
            
            # 返回基础结果
            result = {
                'market_files': len(market_files),
                'stock_files': len(stock_files),
                'market_data': len(market_data),
                'stock_data': 0,
                'top_stocks': {},
                'stock_analysis': {},
                'report_path': output_path,
                'file_preview': self.get_file_preview(market_files + stock_files)
            }
            self.analysis_finished.emit(result)
            
        except Exception as e:
            self.analysis_error.emit(f"生成基础报告失败: {str(e)}")
    
    def generate_basic_analysis_report(self, market_data, stock_data, market_files, stock_files):
        """生成基础分析报告（当无法计算收益率时）"""
        try:
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            output_path = os.path.join(desktop_path, f"股票分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            
            doc = Document()
            doc.add_heading('股票数据分析报告', 0)
            
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")
            
            doc.add_heading('数据扫描结果', level=1)
            doc.add_paragraph(f"• 大盘文件数量: {len(market_files)}")
            doc.add_paragraph(f"• 个股文件数量: {len(stock_files)}")
            doc.add_paragraph(f"• 成功处理股票数量: {len(stock_data)}")
            doc.add_paragraph("")
            
            doc.add_heading('分析说明', level=1)
            doc.add_paragraph("本次分析成功读取了数据文件，但无法计算股票收益率。")
            doc.add_paragraph("可能的原因包括:")
            doc.add_paragraph("• 数据中不包含标准的价格列（如收盘价、价格等）")
            doc.add_paragraph("• 价格数据格式不正确")
            doc.add_paragraph("• 数据量不足")
            doc.add_paragraph("")
            
            doc.add_heading('处理建议', level=1)
            doc.add_paragraph("1. 检查数据文件是否包含价格信息")
            doc.add_paragraph("2. 确认价格列为数值格式")
            doc.add_paragraph("3. 确保有足够的历史数据")
            doc.add_paragraph("4. 查看详细日志了解具体问题")
            
            doc.save(output_path)
            
            self.log_message.emit(f"📄 生成基础分析报告: {output_path}")
            
            # 返回结果
            result = {
                'market_files': len(market_files),
                'stock_files': len(stock_files),
                'market_data': len(market_data),
                'stock_data': len(stock_data),
                'top_stocks': {},
                'stock_analysis': {},
                'report_path': output_path,
                'file_preview': self.get_file_preview(market_files + stock_files)
            }
            self.analysis_finished.emit(result)
            
        except Exception as e:
            self.analysis_error.emit(f"生成基础分析报告失败: {str(e)}")
    
    def generate_basic_market_analysis(self, market_data):
        """生成基础大盘分析"""
        if not market_data:
            return "暂无大盘数据可供分析"
        
        analysis = "【基础大盘分析】\n\n"
        analysis += f"分析时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        analysis += f"数据文件数量: {len(market_data)}\n"
        
        # 合并所有数据
        try:
            combined_data = pd.concat(market_data, ignore_index=True)
            analysis += f"总数据条数: {len(combined_data)}\n"
            analysis += f"数据时间范围: {combined_data['统计日期'].min()} 至 {combined_data['统计日期'].max()}\n"
            
            # 查找价格列进行简单分析
            price_cols = ['收盘', '收盘价', 'close', 'Close']
            for col in price_cols:
                if col in combined_data.columns:
                    prices = pd.to_numeric(combined_data[col], errors='coerce').dropna()
                    if len(prices) > 0:
                        analysis += f"\n价格分析 ({col}):\n"
                        analysis += f"- 平均值: {prices.mean():.2f}\n"
                        analysis += f"- 最高值: {prices.max():.2f}\n"
                        analysis += f"- 最低值: {prices.min():.2f}\n"
                        analysis += f"- 标准差: {prices.std():.2f}\n"
                        break
        except Exception as e:
            analysis += f"\n数据合并分析时出错: {str(e)}"
        
        return analysis
    
    def generate_basic_stock_analysis(self, top_stocks):
        """生成基础个股分析"""
        analysis = {}
        
        for code, info in top_stocks.items():
            stock_analysis = f"【股票 {code} 基础分析】\n\n"
            stock_analysis += f"累计涨幅: {info['total_return']:.2f}%\n"
            stock_analysis += f"起始价格: {info['start_price']:.2f}\n"
            stock_analysis += f"当前价格: {info['end_price']:.2f}\n"
            stock_analysis += f"数据点数: {info['data_points']}\n"
            
            # 简单的投资建议
            if info['total_return'] > 20:
                stock_analysis += "\n投资建议: 表现优秀，可考虑持有或适量加仓"
            elif info['total_return'] > 0:
                stock_analysis += "\n投资建议: 表现良好，可继续观察"
            else:
                stock_analysis += "\n投资建议: 表现一般，建议谨慎操作"
            
            analysis[code] = stock_analysis
        
        return analysis

    def create_word_report(self, top_stocks, market_analysis, stock_analysis, output_path):
        """创建Word报告"""
        doc = Document()
        
        # 标题
        title = doc.add_heading('股票分析报告', 0)
        title.alignment = 1
        
        # 报告信息
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"分析数据: 历史股票数据")
        doc.add_paragraph(f"推荐股票数量: {len(top_stocks)} 只")
        doc.add_paragraph("")
        
        # 大盘分析
        doc.add_heading('一、大盘分析', level=1)
        market_para = doc.add_paragraph(market_analysis)
        
        # 个股推荐
        doc.add_heading('二、推荐股票列表', level=1)
        doc.add_paragraph("基于历史表现分析，推荐以下股票：")
        
        # 创建表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '排名'
        hdr_cells[1].text = '股票代码'
        hdr_cells[2].text = '累计涨幅(%)'
        hdr_cells[3].text = '起始价格'
        hdr_cells[4].text = '当前价格'
        
        # 填充数据
        for i, (code, info) in enumerate(top_stocks.items(), 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = code
            row_cells[2].text = f"{info['total_return']:.2f}%"
            row_cells[3].text = f"{info['start_price']:.2f}"
            row_cells[4].text = f"{info['end_price']:.2f}"
        
        doc.add_paragraph("")
        
        # 个股详细分析
        if stock_analysis:
            doc.add_heading('三、个股详细分析', level=1)
            for code, analysis in stock_analysis.items():
                doc.add_heading(f'股票 {code} 分析', level=2)
                doc.add_paragraph(analysis)
                doc.add_paragraph("")
        
        # 保存文档
        doc.save(output_path)


class StockAnalysisTool(QMainWindow):
    def __init__(self):
        super().__init__()
        self.API_KEY = "sk-2df6ea0568774004950cd5eb2e2adc8a"  # 使用你提供的API密钥
        self.is_analyzing = False
        self.analysis_thread = None
        
        self.init_ui()
    
    def init_ui(self):
        """初始化用户界面"""
        self.setWindowTitle("股票智能分析系统 - 带排错功能")
        self.setGeometry(100, 100, 1400, 900)
        
        # 设置样式
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QGroupBox {
                font-weight: bold;
                border: 2px solid #cccccc;
                border-radius: 5px;
                margin-top: 1ex;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
            }
            QPushButton {
                background-color: #4CAF50;
                border: none;
                color: white;
                padding: 8px 16px;
                text-align: center;
                text-decoration: none;
                font-size: 14px;
                margin: 4px 2px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
            QPushButton:disabled {
                background-color: #cccccc;
            }
            QPushButton#danger {
                background-color: #f44336;
            }
            QPushButton#danger:hover {
                background-color: #da190b;
            }
            QPushButton#primary {
                background-color: #2196F3;
            }
            QPushButton#primary:hover {
                background-color: #0b7dda;
            }
            QTextEdit {
                border: 1px solid #cccccc;
                border-radius: 3px;
                padding: 5px;
                font-family: 'Courier New';
                font-size: 10px;
            }
            QProgressBar {
                border: 1px solid #cccccc;
                border-radius: 3px;
                text-align: center;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                width: 10px;
            }
            QTableWidget {
                font-size: 11px;
            }
        """)
        
        # 创建中心部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        layout = QVBoxLayout(central_widget)
        
        # 标题
        title_label = QLabel("📈 股票智能分析系统 - 智能排错版")
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: bold;
                color: #2c3e50;
                padding: 20px;
                background-color: #e3f2fd;
                border-radius: 10px;
                margin: 10px;
            }
        """)
        layout.addWidget(title_label)
        
        # 窗口控制选项
        control_layout = QHBoxLayout()
        self.always_on_top_check = QCheckBox("窗口始终置顶")
        self.always_on_top_check.setChecked(False)
        self.always_on_top_check.toggled.connect(self.toggle_always_on_top)
        control_layout.addWidget(self.always_on_top_check)
        control_layout.addStretch()
        layout.addLayout(control_layout)
        
        # 创建选项卡
        tab_widget = QTabWidget()
        layout.addWidget(tab_widget)
        
        # 分析选项卡
        analysis_tab = QWidget()
        self.setup_analysis_tab(analysis_tab)
        tab_widget.addTab(analysis_tab, "📊 数据与分析")
        
        # 文件预览选项卡
        preview_tab = QWidget()
        self.setup_preview_tab(preview_tab)
        tab_widget.addTab(preview_tab, "👀 文件预览")
        
        # 结果选项卡
        result_tab = QWidget()
        self.setup_result_tab(result_tab)
        tab_widget.addTab(result_tab, "📈 分析结果")
        
        # 日志选项卡
        log_tab = QWidget()
        self.setup_log_tab(log_tab)
        tab_widget.addTab(log_tab, "📝 详细日志")
    
    def toggle_always_on_top(self, checked):
        """切换窗口置顶状态"""
        if checked:
            self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)
        else:
            self.setWindowFlags(self.windowFlags() & ~Qt.WindowStaysOnTopHint)
        self.show()
    
    def setup_analysis_tab(self, parent):
        """设置分析选项卡"""
        layout = QVBoxLayout(parent)
        
        # API状态
        api_group = QGroupBox("🔑 DeepSeek API状态")
        api_layout = QVBoxLayout(api_group)
        api_status = QLabel(f"API密钥: {self.API_KEY[:8]}...{self.API_KEY[-4:]} (已配置)")
        api_status.setStyleSheet("color: #4CAF50; font-weight: bold; padding: 5px;")
        api_layout.addWidget(api_status)
        
        # 排错功能说明
        error_handling_label = QLabel("💡 排错功能已启用: 自动处理数据格式问题，跳过错误文件")
        error_handling_label.setStyleSheet("color: #FF9800; font-weight: bold; padding: 5px; background-color: #FFF3E0; border-radius: 3px;")
        api_layout.addWidget(error_handling_label)
        
        layout.addWidget(api_group)
        
        # 数据管理组
        data_group = QGroupBox("📁 数据管理")
        data_layout = QVBoxLayout(data_group)
        
        # 文件夹选择
        folder_layout = QHBoxLayout()
        folder_layout.addWidget(QLabel("数据文件夹:"))
        
        self.folder_edit = QLineEdit()
        self.folder_edit.setPlaceholderText("请选择包含股票数据的文件夹...")
        folder_layout.addWidget(self.folder_edit)
        
        self.browse_btn = QPushButton("浏览文件夹")
        self.browse_btn.clicked.connect(self.browse_folder)
        folder_layout.addWidget(self.browse_btn)
        
        data_layout.addLayout(folder_layout)
        
        # 数据信息
        self.data_info_label = QLabel("请选择包含股票数据的文件夹")
        self.data_info_label.setStyleSheet("color: #666666;")
        data_layout.addWidget(self.data_info_label)
        
        layout.addWidget(data_group)
        
        # 分析选项组
        analysis_group = QGroupBox("🔍 分析选项")
        analysis_layout = QVBoxLayout(analysis_group)
        
        # 分析类型
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel("分析类型:"))
        
        self.analysis_type_group = QButtonGroup()
        self.comprehensive_radio = QRadioButton("全面分析（大盘+个股）")
        self.market_only_radio = QRadioButton("仅大盘分析")
        self.stock_only_radio = QRadioButton("仅个股分析")
        
        self.comprehensive_radio.setChecked(True)
        
        self.analysis_type_group.addButton(self.comprehensive_radio)
        self.analysis_type_group.addButton(self.market_only_radio)
        self.analysis_type_group.addButton(self.stock_only_radio)
        
        type_layout.addWidget(self.comprehensive_radio)
        type_layout.addWidget(self.market_only_radio)
        type_layout.addWidget(self.stock_only_radio)
        type_layout.addStretch()
        
        analysis_layout.addLayout(type_layout)
        
        # 预测设置
        predict_layout = QHBoxLayout()
        predict_layout.addWidget(QLabel("预测天数:"))
        
        self.predict_days_edit = QLineEdit("7")
        self.predict_days_edit.setFixedWidth(50)
        predict_layout.addWidget(self.predict_days_edit)
        predict_layout.addWidget(QLabel("天"))
        predict_layout.addStretch()
        
        analysis_layout.addLayout(predict_layout)
        
        layout.addWidget(analysis_group)
        
        # 进度组
        progress_group = QGroupBox("⏳ 分析进度")
        progress_layout = QVBoxLayout(progress_group)
        
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        progress_layout.addWidget(self.progress_bar)
        
        self.progress_label = QLabel("等待开始分析...")
        self.progress_label.setStyleSheet("font-weight: bold; color: #2c3e50;")
        progress_layout.addWidget(self.progress_label)
        
        self.detail_label = QLabel("")
        self.detail_label.setStyleSheet("color: #666666;")
        progress_layout.addWidget(self.detail_label)
        
        layout.addWidget(progress_group)
        
        # 控制按钮
        button_layout = QHBoxLayout()
        
        self.analyze_btn = QPushButton("开始智能分析（带排错）")
        self.analyze_btn.clicked.connect(self.start_analysis)
        self.analyze_btn.setStyleSheet("font-size: 16px; padding: 10px 20px;")
        self.analyze_btn.setObjectName("primary")
        button_layout.addWidget(self.analyze_btn)
        
        self.reset_btn = QPushButton("清空数据")
        self.reset_btn.clicked.connect(self.reset_all)
        self.reset_btn.setObjectName("danger")
        button_layout.addWidget(self.reset_btn)
        
        button_layout.addStretch()
        
        layout.addLayout(button_layout)
    
    def setup_preview_tab(self, parent):
        """设置文件预览选项卡"""
        layout = QVBoxLayout(parent)
        
        # 预览说明
        preview_label = QLabel("👀 文件数据预览（扫描后显示前3个文件的信息）")
        preview_label.setStyleSheet("font-size: 16px; font-weight: bold; padding: 10px;")
        layout.addWidget(preview_label)
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setPlaceholderText("请先选择数据文件夹并扫描文件...")
        layout.addWidget(self.preview_text)
    
    def setup_result_tab(self, parent):
        """设置结果选项卡"""
        layout = QVBoxLayout(parent)
        
        # 结果表格
        result_label = QLabel("📊 推荐股票列表")
        result_label.setStyleSheet("font-size: 16px; font-weight: bold; padding: 10px;")
        layout.addWidget(result_label)
        
        self.result_table = QTableWidget()
        self.result_table.setColumnCount(5)
        self.result_table.setHorizontalHeaderLabels(['排名', '股票代码', '累计涨幅(%)', '起始价格', '当前价格'])
        self.result_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        layout.addWidget(self.result_table)
        
        # 报告信息
        report_group = QGroupBox("📄 分析报告")
        report_layout = QVBoxLayout(report_group)
        self.report_info_label = QLabel("暂无分析报告")
        self.report_info_label.setStyleSheet("color: #666666; padding: 10px;")
        self.report_info_label.setWordWrap(True)
        report_layout.addWidget(self.report_info_label)
        layout.addWidget(report_group)
    
    def setup_log_tab(self, parent):
        """设置日志选项卡"""
        layout = QVBoxLayout(parent)
        
        log_label = QLabel("📝 详细分析日志")
        log_label.setStyleSheet("font-size: 16px; font-weight: bold; padding: 10px;")
        layout.addWidget(log_label)
        
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        layout.addWidget(self.log_text)
    
    def browse_folder(self):
        """浏览文件夹"""
        folder_path = QFileDialog.getExistingDirectory(self, "选择股票数据文件夹")
        if folder_path:
            self.folder_edit.setText(folder_path)
            self.log_message(f"✅ 已选择数据文件夹: {folder_path}")
            self.scan_data_files(folder_path)
    
    def scan_data_files(self, folder_path):
        """扫描数据文件"""
        try:
            market_files = []
            stock_files = []
            
            for pattern in ["*.xls", "*.xlsx"]:
                search_pattern = os.path.join(folder_path, "**", pattern)
                for file_path in glob.glob(search_pattern, recursive=True):
                    filename = os.path.basename(file_path)
                    if '板块' in filename:
                        market_files.append(file_path)
                    else:
                        stock_files.append(file_path)
            
            file_count = len(market_files) + len(stock_files)
            self.data_info_label.setText(f"找到 {len(market_files)} 个大盘文件, {len(stock_files)} 个个股文件，共 {file_count} 个文件")
            self.log_message(f"📁 找到 {len(market_files)} 个大盘文件, {len(stock_files)} 个个股文件")
            
            # 更新文件预览
            self.update_file_preview(market_files + stock_files)
            
        except Exception as e:
            self.log_message(f"❌ 扫描文件时出错: {str(e)}")
    
    def update_file_preview(self, file_paths):
        """更新文件预览"""
        preview_text = "文件预览信息:\n\n"
        
        if not file_paths:
            preview_text += "未找到任何文件"
            self.preview_text.setPlainText(preview_text)
            return
        
        for i, file_path in enumerate(file_paths[:3]):  # 只预览前3个文件
            preview_text += f"=== 文件 {i+1}: {os.path.basename(file_path)} ===\n"
            
            try:
                # 尝试读取文件
                for engine in [None, 'xlrd', 'openpyxl']:
                    try:
                        df = pd.read_excel(file_path, engine=engine, nrows=10)  # 只读取前10行
                        if df is not None and not df.empty:
                            preview_text += f"✅ 读取成功 (引擎: {engine})\n"
                            preview_text += f"数据形状: {df.shape}\n"
                            preview_text += f"列名: {list(df.columns)}\n"
                            preview_text += "前10行数据:\n"
                            preview_text += df.head(10).to_string() + "\n\n"
                            break
                    except:
                        continue
                else:
                    preview_text += "❌ 读取失败\n\n"
            except Exception as e:
                preview_text += f"❌ 读取失败: {str(e)}\n\n"
        
        self.preview_text.setPlainText(preview_text)
    
    def log_message(self, message):
        """添加日志消息"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        self.log_text.append(log_entry)
    
    def update_progress(self, value, message, detail):
        """更新进度"""
        self.progress_bar.setValue(value)
        self.progress_label.setText(message)
        if detail:
            self.detail_label.setText(detail)
    
    def start_analysis(self):
        """开始分析"""
        if self.is_analyzing:
            return
        
        folder_path = self.folder_edit.text().strip()
        if not folder_path or not os.path.exists(folder_path):
            QMessageBox.warning(self, "错误", "请选择有效的数据文件夹")
            return
        
        try:
            predict_days = int(self.predict_days_edit.text().strip())
        except:
            predict_days = 7
        
        # 获取分析类型
        if self.comprehensive_radio.isChecked():
            analysis_type = "comprehensive"
        elif self.market_only_radio.isChecked():
            analysis_type = "market_only"
        else:
            analysis_type = "stock_only"
        
        self.is_analyzing = True
        self.analyze_btn.setEnabled(False)
        self.reset_btn.setEnabled(False)
        
        # 创建分析线程
        self.analysis_thread = AnalysisThread(folder_path, analysis_type, predict_days, self.API_KEY)
        self.analysis_thread.progress_updated.connect(self.update_progress)
        self.analysis_thread.log_message.connect(self.log_message)
        self.analysis_thread.analysis_finished.connect(self.on_analysis_finished)
        self.analysis_thread.analysis_error.connect(self.on_analysis_error)
        self.analysis_thread.start()
    
    def on_analysis_finished(self, result):
        """分析完成"""
        self.is_analyzing = False
        self.analyze_btn.setEnabled(True)
        self.reset_btn.setEnabled(True)
        
        # 更新结果表格
        self.update_result_table(result['top_stocks'])
        
        # 显示报告信息
        report_info = f"✅ 分析完成!\n\n"
        report_info += f"📊 处理统计:\n"
        report_info += f"• 大盘文件: {result['market_files']} 个\n"
        report_info += f"• 个股文件: {result['stock_files']} 个\n"
        report_info += f"• 成功处理: {result['stock_data']} 只股票\n"
        report_info += f"• 推荐股票: {len(result['top_stocks'])} 只\n\n"
        report_info += f"📄 报告位置:\n{result['report_path']}"
        
        self.report_info_label.setText(report_info)
        
        # 显示完成消息
        result_msg = f"分析完成！\n\n"
        result_msg += f"• 处理大盘文件: {result['market_files']} 个\n"
        result_msg += f"• 处理个股文件: {result['stock_files']} 个\n"
        result_msg += f"• 推荐股票: {len(result['top_stocks'])} 只\n"
        result_msg += f"• 报告位置: {result['report_path']}"
        
        QMessageBox.information(self, "分析完成", result_msg)
    
    def on_analysis_error(self, error_msg):
        """分析错误"""
        self.is_analyzing = False
        self.analyze_btn.setEnabled(True)
        self.reset_btn.setEnabled(True)
        
        QMessageBox.critical(self, "分析错误", f"分析过程中出错:\n{error_msg}")
        self.log_message(f"❌ 分析错误: {error_msg}")
    
    def update_result_table(self, top_stocks):
        """更新结果表格"""
        self.result_table.setRowCount(len(top_stocks))
        
        for i, (code, info) in enumerate(top_stocks.items()):
            self.result_table.setItem(i, 0, QTableWidgetItem(str(i + 1)))
            self.result_table.setItem(i, 1, QTableWidgetItem(code))
            self.result_table.setItem(i, 2, QTableWidgetItem(f"{info['total_return']:.2f}%"))
            self.result_table.setItem(i, 3, QTableWidgetItem(f"{info['start_price']:.2f}"))
            self.result_table.setItem(i, 4, QTableWidgetItem(f"{info['end_price']:.2f}"))
    
    def reset_all(self):
        """重置所有数据"""
        reply = QMessageBox.question(self, "确认", "确定要清空所有数据吗？",
                                   QMessageBox.Yes | QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            self.folder_edit.clear()
            self.data_info_label.setText("请选择包含股票数据的文件夹")
            self.progress_bar.setValue(0)
            self.progress_label.setText("等待开始分析...")
            self.detail_label.clear()
            self.log_text.clear()
            self.preview_text.clear()
            self.result_table.setRowCount(0)
            self.report_info_label.setText("暂无分析报告")
            self.log_message("系统已重置")


def main():
    # 创建应用实例
    app = QApplication(sys.argv)
    app.setApplicationName("股票智能分析系统 - 智能排错版")
    
    # 创建并显示主窗口
    window = StockAnalysisTool()
    window.show()
    
    # 执行应用
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()